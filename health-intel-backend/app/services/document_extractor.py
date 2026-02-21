"""
Document Extraction Service
Extracts plain text from:
  - PDF files  (using pdfplumber — no OCR needed for standard medical reports)
  - DOCX files (using python-docx)

The extracted text is returned for use as a MedicalRecord summary,
injected into assessment context, or displayed back to the user.
"""
import io
from pathlib import Path


def extract_text_from_pdf(content: bytes) -> str:
    """
    Extract text from a PDF file byte string.
    Uses pdfplumber for reliable text extraction across multi-page documents.

    Returns:
        Extracted text as a single string (pages separated by newlines).
    """
    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError(
            "pdfplumber is not installed. Run: pip install pdfplumber"
        )

    text_parts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_parts.append(f"--- Page {i + 1} ---\n{page_text.strip()}")

    if not text_parts:
        return ""
    return "\n\n".join(text_parts)


def extract_text_from_docx(content: bytes) -> str:
    """
    Extract text from a .docx file byte string.
    Uses python-docx to iterate through paragraphs and tables.

    Returns:
        Extracted text as a single string.
    """
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    doc = Document(io.BytesIO(content))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Also extract text from tables (lab results are often tabular)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def extract_text(filename: str, content: bytes) -> str:
    """
    Route extraction to the right handler based on file extension.

    Args:
        filename: Original filename (used to detect format)
        content:  Raw bytes of the uploaded file

    Returns:
        Extracted text string

    Raises:
        ValueError: If the file type is not supported
        RuntimeError: If a required library is missing
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return extract_text_from_pdf(content)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(content)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            "Only .pdf and .docx files are supported."
        )


def summarize_extracted_text(text: str, max_chars: int = 3000) -> str:
    """
    Truncate extracted text to a manageable length for DB storage.
    Preserves the beginning (patient info) and end (conclusions/findings).

    For very long reports (e.g. full blood panels), we keep the first 2/3
    and the last 1/3 of the allowed budget so key findings are included.
    """
    if len(text) <= max_chars:
        return text

    head_budget = int(max_chars * 0.65)
    tail_budget = max_chars - head_budget - 50  # 50 chars for separator

    head = text[:head_budget]
    tail = text[-tail_budget:]
    return f"{head}\n\n[... content truncated ...]\n\n{tail}"
